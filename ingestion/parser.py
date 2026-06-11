# ingestion/parser.py
from pathlib import Path
from dataclasses import dataclass, field
from loguru import logger
import pandas as pd

@dataclass
class ParsedDocument:
    source: str
    file_type: str
    text: str = ""
    tables: list[pd.DataFrame] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

class DocumentParser:
    SUPPORTED = {".pdf", ".docx", ".csv", ".xlsx", ".xls", ".txt"}

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()
        logger.info(f"Parsing {path.name}")

        if suffix == ".pdf":       return self._parse_pdf(path)
        elif suffix == ".docx":    return self._parse_docx(path)
        elif suffix in {".csv", ".xlsx", ".xls"}: return self._parse_tabular(path)
        elif suffix == ".txt":     return self._parse_txt(path)
        else: raise ValueError(f"Unsupported: {suffix}")

    def _parse_pdf(self, path):
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
        return ParsedDocument(
            source=str(path), file_type="pdf", text=text,
            metadata={"pages": len(reader.pages), "filename": path.name}
        )

    def _parse_docx(self, path):
        from docx import Document
        doc = Document(str(path))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        tables = []
        for table in doc.tables:
            rows = [[c.text for c in r.cells] for r in table.rows]
            if rows:
                tables.append(pd.DataFrame(rows[1:], columns=rows[0]))
        return ParsedDocument(
            source=str(path), file_type="docx", text=text,
            tables=tables, metadata={"filename": path.name}
        )

    def _parse_tabular(self, path):
        df = pd.read_csv(path) if path.suffix == ".csv" else pd.read_excel(path)
        text = f"Table: {path.name}\n{df.to_string(index=False)}"
        return ParsedDocument(
            source=str(path), file_type=path.suffix.lstrip("."),
            text=text, tables=[df],
            metadata={"filename": path.name, "rows": len(df), "cols": len(df.columns)}
        )

    def _parse_txt(self, path):
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(source=str(path), file_type="txt",
                              text=text, metadata={"filename": path.name})